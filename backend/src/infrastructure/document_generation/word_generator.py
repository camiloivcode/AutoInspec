from typing import Dict, Any
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ...domain.services.document_generator import DocumentGeneratorService


class WordGeneratorService(DocumentGeneratorService):
    async def generate_docx(self, template_content: str, variables: Dict[str, Any], output_path: str) -> str:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        sections = template_content.split('\n\n')
        for section in sections:
            section = section.strip()
            if not section:
                continue

            lines = section.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                resolved = self._resolve_variables(line, variables)

                if line.startswith('# '):
                    heading = doc.add_heading(resolved[2:], level=1)
                elif line.startswith('## '):
                    heading = doc.add_heading(resolved[3:], level=2)
                elif line.startswith('### '):
                    heading = doc.add_heading(resolved[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(resolved[2:], style='List Bullet')
                elif '|' in line and line.count('|') >= 3:
                    pass
                elif line.startswith('|'):
                    pass
                else:
                    p = doc.add_paragraph(resolved)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        table_data = self._extract_table(template_content, variables)
        if table_data:
            rows = table_data
            if len(rows) > 1:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, row in enumerate(rows):
                    for j, cell_text in enumerate(row):
                        table.cell(i, j).text = str(cell_text)

        doc.save(output_path)
        return output_path

    async def generate_pdf(self, template_content: str, variables: Dict[str, Any], output_path: str) -> str:
        return await self.generate_docx(template_content, variables, output_path)

    async def get_preview(self, template_content: str, variables: Dict[str, Any]) -> str:
        text = template_content
        for key, value in variables.items():
            text = text.replace(f"{{{{{key}}}}}", str(value))
        return text[:500]

    def _resolve_variables(self, text: str, variables: Dict[str, Any]) -> str:
        result = text
        for key, value in variables.items():
            result = result.replace(f"{{{{{key}}}}}", str(value) if value is not None else "")
        return result

    def _extract_table(self, template_content: str, variables: Dict[str, Any]) -> list:
        rows = []
        for line in template_content.split('\n'):
            line = line.strip()
            if line.startswith('|') and line.endswith('|'):
                cells = [c.strip() for c in line.strip('|').split('|')]
                resolved_row = [self._resolve_variables(c, variables) for c in cells]
                if not all(c == '---' for c in resolved_row):
                    rows.append(resolved_row)
        return rows
