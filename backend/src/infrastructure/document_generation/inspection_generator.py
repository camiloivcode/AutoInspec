import os
import re
import logging
import tempfile
from typing import Dict, List

from PIL import Image
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

PHOTO_MARGIN = Cm(0)

EMU_PER_INCH = 914400
DEFAULT_DPI = 96
PX_TO_EMU = EMU_PER_INCH // DEFAULT_DPI

MAX_PDF_PX = 1600


class InspectionPDFGenerator:
    def __init__(self):
        self._orient_cache: dict[str, tuple[int, int]] = {}

    def _get_oriented_size(self, img_path: str) -> tuple[int, int]:
        cached = self._orient_cache.get(img_path)
        if cached:
            return cached
        with Image.open(img_path) as img:
            w, h = img.size
            try:
                orientation = img.getexif().get(0x0112, 1)
            except Exception:
                orientation = 1
            if orientation in (5, 6, 7, 8):
                w, h = h, w
            self._orient_cache[img_path] = (w, h)
            return w, h

    def _resize_for_pdf(self, img_path: str) -> str:
        with Image.open(img_path) as img:
            try:
                orientation = img.getexif().get(0x0112, 1)
            except Exception:
                orientation = 1
            if orientation in (5, 6, 7, 8):
                img = img.transpose(Image.Transpose.ROTATE_270 if orientation == 6 else (
                    Image.Transpose.ROTATE_90 if orientation == 8 else (
                        Image.Transpose.ROTATE_180 if orientation == 3 else Image.Transpose.FLIP_LEFT_RIGHT)))
            w, h = img.size
            longest = max(w, h)
            if longest > MAX_PDF_PX:
                scale = MAX_PDF_PX / longest
                new_w = int(w * scale)
                new_h = int(h * scale)
                img = img.resize((new_w, new_h), Image.LANCZOS)
            fd, tmp_path = tempfile.mkstemp(suffix='.jpg')
            os.close(fd)
            img.save(tmp_path, 'JPEG', quality=92)
            return tmp_path

    def generate(self, driver_name: str, plate: str, images: Dict[int, List[str]], output_dir: str) -> str:
        doc = Document()

        section = doc.sections[0]
        section.top_margin = PHOTO_MARGIN
        section.bottom_margin = PHOTO_MARGIN
        section.left_margin = PHOTO_MARGIN
        section.right_margin = PHOTO_MARGIN

        avail_w = section.page_width - section.left_margin - section.right_margin
        avail_h = section.page_height - section.top_margin - section.bottom_margin

        tmp_files: list[str] = []
        sorted_positions = sorted(images.keys())
        try:
            for idx, pos in enumerate(sorted_positions):
                for img_idx, img_path in enumerate(images[pos]):
                    if idx > 0 or img_idx > 0:
                        doc.add_section()
                        new_section = doc.sections[-1]
                        new_section.top_margin = PHOTO_MARGIN
                        new_section.bottom_margin = PHOTO_MARGIN
                        new_section.left_margin = PHOTO_MARGIN
                        new_section.right_margin = PHOTO_MARGIN

                    if os.path.exists(img_path):
                        try:
                            resized = self._resize_for_pdf(img_path)
                            if resized != img_path:
                                tmp_files.append(resized)
                            pw, ph = self._get_oriented_size(img_path)
                            native_w = pw * PX_TO_EMU
                            native_h = ph * PX_TO_EMU

                            scale = min(avail_w / native_w, avail_h / native_h, 1.0)
                            draw_w = int(native_w * scale)
                            draw_h = int(native_h * scale)

                            doc.add_picture(resized, width=draw_w, height=draw_h)
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            logger.warning(f"Error inserting image {img_path}: {e}")

            safe_name = re.sub(r'[^\w\s-]', '', driver_name).strip().replace(' ', '_')
            os.makedirs(output_dir, exist_ok=True)
            docx_path = os.path.join(output_dir, f"{safe_name}.docx")
            doc.save(docx_path)

            pdf_path = self._convert_to_pdf(docx_path)
            return pdf_path
        finally:
            for tmp in tmp_files:
                try:
                    os.remove(tmp)
                except Exception:
                    pass
            self._orient_cache.clear()

    def _convert_to_pdf(self, docx_path: str) -> str:
        pdf_path = docx_path.replace(".docx", ".pdf")

        try:
            import subprocess
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(pdf_path), docx_path],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0:
                logger.info(f"PDF generated via LibreOffice: {pdf_path}")
                return pdf_path
            logger.warning(f"LibreOffice failed: {result.stderr}")
        except Exception as e:
            logger.warning(f"LibreOffice not available: {e}")

        logger.warning(f"PDF conversion not available, returning DOCX: {docx_path}")
        return docx_path
